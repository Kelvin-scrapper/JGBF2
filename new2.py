#!/usr/bin/env python3
"""
Title-Enhanced Selective Page PDF→DOCX→Excel Converter
Pre-scans PDFs to find pages with specific subtitles and converts only those pages.
"""

import os
import sys
import multiprocessing
from pathlib import Path
import logging
from typing import Optional, List, Tuple, Dict
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass
import re

try:
    import fitz  # PyMuPDF
    from pdf2docx import Converter
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
except ImportError as e:
    print(f"Missing required packages. Install with:")
    print("pip install PyMuPDF pdf2docx python-docx openpyxl")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


@dataclass
class PageJob:
    """Represents a single-page processing job"""
    pdf_path: Path
    pdf_name: str
    page_number: int  # 0-indexed page number
    subtitle: str     # The subtitle found on this page
    job_id: str       # Unique identifier for this job
    chunks_folder: Path
    table_titles: List[str] = None


def convert_page_worker(job: PageJob):
    """Worker function for converting a single, pre-identified relevant page."""
    page_docx_path = job.chunks_folder / f"{job.pdf_name}_page_{job.page_number}.docx"
    
    try:
        logger.info(f"--> Converting {job.job_id}: Page {job.page_number + 1} ('{job.subtitle}')")
        
        cv = Converter(str(job.pdf_path))
        cv.convert(
            str(page_docx_path),
            start=job.page_number,
            end=job.page_number + 1,
            table_settings={'snap_tolerance': 1.0, 'min_border_width': 0.3, 'join_tolerance': 1.0}
        )
        cv.close()
        
        chunk_doc = Document(str(page_docx_path))
        table_count = len(chunk_doc.tables)
        
        logger.info(f"[OK] {job.job_id}: Completed with {table_count} tables.")
        
        return {
            'job_id': job.job_id,
            'pdf_name': job.pdf_name,
            'page_number': job.page_number,
            'page_path': page_docx_path,
            'subtitle': job.subtitle,
            'table_count': table_count,
            'table_titles': job.table_titles,
            'success': True
        }
        
    except Exception as e:
        logger.error(f"[ERROR] {job.job_id} failed: {e}")
        return {
            'job_id': job.job_id,
            'pdf_name': job.pdf_name,
            'page_number': job.page_number,
            'success': False,
            'error': str(e)
        }


class TitleEnhancedConverter:
    """Selectively converts PDF pages based on subtitles."""
    
    def __init__(self, pdf_folder: str, excel_output_folder: str, max_workers: int = 6):
        self.pdf_folder = Path(pdf_folder)
        self.excel_output_folder = Path(excel_output_folder)
        self.docx_folder = self.excel_output_folder / "converted_docx"
        self.chunks_folder = self.excel_output_folder / "temp_chunks"
        
        self.excel_output_folder.mkdir(exist_ok=True)
        self.docx_folder.mkdir(exist_ok=True)
        self.chunks_folder.mkdir(exist_ok=True)
        
        self.max_workers = max_workers
        
        # Define the subtitles to search for.
        self.target_subtitles = [
            "長期国債先物（現金決済型ミニ）",
            "JGB(10-year) Futures",
            "mini-10-year JGB Futures (Cash-Settled)",
            "mini-20-year JGB Futures",
            "3-Month TONA Futures"
        ]
        
        # Define keywords that identify an "Options" page to be excluded
        self.primary_exclusion_keywords = ["Options on", "オプション"]

        # Sort targets by length (longest first) to prioritize more specific matches.
        self.target_subtitles.sort(key=len, reverse=True)

        self.table_section_titles = [
            "総計・自己合計・委託合計 Total, Proprietary & Brokerage",
            "委託内訳 Breakdown of Brokerage",
            "法人内訳 Breakdown of Institutions",
            "金融機関内訳 Breakdown of Financial Institutions"
        ]
        self.table_keywords = ["総計・自己合計・委託合計", "委託内訳", "法人内訳", "金融機関内訳"]
        self.table_names = ["Table1_Main_Summary", "Table2_Brokerage_Breakdown", "Table3_Institutions_Breakdown", "Table4_Financial_Breakdown"]
        
        logger.info(f"[INFO] Selective Page Converter configured:")
        logger.info(f"  • PDF input folder: {self.pdf_folder}")
        logger.info(f"  • Excel output folder: {self.excel_output_folder}")

    def extract_table_titles_from_text(self, text: str) -> List[str]:
        found_titles = []
        for keyword, full_title in zip(self.table_keywords, self.table_section_titles):
            if keyword in text:
                found_titles.append(full_title)
        while len(found_titles) < len(self.table_keywords):
            found_titles.append(f"Table Title {len(found_titles) + 1} (Not Found)")
        return found_titles

    def extract_table_titles_for_page(self, pdf_path: Path, page_number: int) -> List[str]:
        try:
            doc = fitz.open(str(pdf_path))
            page = doc[page_number]
            text = page.get_text("text")
            doc.close()
            return self.extract_table_titles_from_text(text)
        except Exception as e:
            logger.error(f"Could not extract table titles from page {page_number + 1}: {e}")
            return self.table_section_titles

    def get_relevant_pages_and_subtitles(self, pdf_path: Path) -> Dict[int, str]:
        """Scans a PDF and captures the full line of text for the subtitle."""
        relevant_pages = {}
        try:
            doc = fitz.open(str(pdf_path))
            logger.info(f"[SCAN] Precisely scanning {pdf_path.name} ({len(doc)} pages)...")
            
            for i, page in enumerate(doc):
                page_match_found = False
                lines = page.get_text("text").splitlines()
                
                for line in lines:
                    if not line.strip():
                        continue

                    is_excluded_line = any(ex_keyword in line for ex_keyword in self.primary_exclusion_keywords)
                    if is_excluded_line:
                        continue
                        
                    for subtitle_keyword in self.target_subtitles:
                        if subtitle_keyword in line:
                            full_subtitle = line.strip()
                            logger.info(f"  > Page {i + 1}: MATCH for '{subtitle_keyword}' -> Capturing full title: '{full_subtitle}'")
                            relevant_pages[i] = full_subtitle
                            page_match_found = True
                            break
                    
                    if page_match_found:
                        break
            doc.close()
        except Exception as e:
            logger.error(f"Could not scan PDF {pdf_path.name}: {e}")
        return relevant_pages

    def process_all_files(self):
        if not self.pdf_folder.exists():
            logger.error(f"PDF input folder '{self.pdf_folder}' does not exist!")
            return
        pdf_files = list(self.pdf_folder.glob("*.pdf"))
        if not pdf_files:
            logger.error(f"No PDF files found in '{self.pdf_folder}'")
            return
        print(f"[INFO] Found {len(pdf_files)} PDF files to process in '{self.pdf_folder}'.")
        self.process_pdfs_selectively(pdf_files)
        self.cleanup_temp_files()

    def process_pdfs_selectively(self, pdf_files: List[Path]):
        start_time = time.time()
        all_page_jobs = []
        for pdf_file in pdf_files:
            relevant_pages = self.get_relevant_pages_and_subtitles(pdf_file)
            if not relevant_pages:
                logger.warning(f"[WARN] No relevant pages found in {pdf_file.name}. Skipping.")
                continue
            pdf_name = pdf_file.stem
            for page_num, subtitle in relevant_pages.items():
                table_titles = self.extract_table_titles_for_page(pdf_file, page_num)
                job = PageJob(
                    pdf_path=pdf_file,
                    pdf_name=pdf_name,
                    page_number=page_num,
                    subtitle=subtitle,
                    job_id=f"{pdf_name}-P{page_num+1}",
                    chunks_folder=self.chunks_folder,
                    table_titles=table_titles
                )
                all_page_jobs.append(job)
        if not all_page_jobs:
            logger.info("No relevant pages to process across all files.")
            return
        total_pages_to_process = len(all_page_jobs)
        logger.info(f"--> Starting parallel conversion of {total_pages_to_process} relevant pages...")
        page_results = []
        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_job = {executor.submit(convert_page_worker, job): job for job in all_page_jobs}
            for future in as_completed(future_to_job):
                result = future.result()
                page_results.append(result)
        parallel_time = time.time() - start_time
        logger.info(f"--> Parallel conversion finished in {parallel_time:.2f} seconds.")
        pdf_results = {}
        for result in page_results:
            if result['success']:
                pdf_name = result['pdf_name']
                if pdf_name not in pdf_results:
                    pdf_results[pdf_name] = []
                pdf_results[pdf_name].append(result)
        for pdf_name, results in pdf_results.items():
            pdf_start_time = time.time()
            logger.info(f"--> Processing results for {pdf_name}...")
            results.sort(key=lambda x: x['page_number'])
            page_paths = [Path(r['page_path']) for r in results]
            combined_docx_path = self.docx_folder / f"{pdf_name}.docx"
            if self.combine_pages(page_paths, combined_docx_path):
                self.convert_docx_to_excel(combined_docx_path, results)
            pdf_time = time.time() - pdf_start_time
            print(f"\n[SUCCESS] {pdf_name} COMPLETED in {pdf_time:.2f}s")
            print(f"  - Converted {len(results)} relevant pages.")
            print(f"  - Excel output: {self.excel_output_folder / (pdf_name + '.xlsx')}")

    def combine_pages(self, page_paths: List[Path], output_path: Path) -> bool:
        if not page_paths: return False
        try:
            combined_doc = Document(str(page_paths[0]))
            for path in page_paths[1:]:
                sub_doc = Document(str(path))
                for element in sub_doc.element.body:
                    combined_doc.element.body.append(element)
            combined_doc.save(str(output_path))
            logger.info(f"[OK] Combined {len(page_paths)} page(s) into {output_path.name}")
            return True
        except Exception as e:
            logger.error(f"Failed to combine pages into {output_path.name}: {e}")
            return False

    def convert_docx_to_excel(self, docx_path: Path, page_results: List[Dict]):
        try:
            doc = Document(docx_path)
            total_tables = len(doc.tables)
            if total_tables == 0:
                logger.warning(f"[WARN] No tables found in {docx_path.name}")
                return
            table_to_subtitle_map = {}
            table_to_table_title_map = {}
            current_table_index = 0
            for result in page_results:
                num_tables_on_page = result['table_count']
                page_table_titles = result.get('table_titles', self.table_section_titles)
                for i in range(num_tables_on_page):
                    table_to_subtitle_map[current_table_index] = result['subtitle']
                    table_pos_on_page = i % 4
                    if table_pos_on_page < len(page_table_titles):
                        table_to_table_title_map[current_table_index] = page_table_titles[table_pos_on_page]
                    else:
                        table_to_table_title_map[current_table_index] = f"Table Title {table_pos_on_page + 1}"
                    current_table_index += 1
            wb = Workbook()
            wb.remove(wb.active)
            self.create_summary_sheet(wb, docx_path.name, len(page_results), total_tables)
            for i, table in enumerate(doc.tables):
                page_number_guess = (i // 4) + 1
                table_position = i % 4
                table_name = self.table_names[table_position]
                sheet_name = f"P{page_number_guess}_{table_name[:20]}"
                ws = wb.create_sheet(title=sheet_name)
                subtitle = table_to_subtitle_map.get(i, "Subtitle Not Found")
                table_title = table_to_table_title_map.get(i, "Table Title Not Found")
                main_title = "Trading by Type of Investors"
                self.copy_table_to_sheet_with_enhanced_titles(table, ws, main_title, subtitle, table_title)
            excel_path = self.excel_output_folder / f"{docx_path.stem}.xlsx"
            wb.save(excel_path)
        except Exception as e:
            logger.error(f"[ERROR] Error processing {docx_path.name} to Excel: {e}")

    def copy_table_to_sheet_with_enhanced_titles(self, table, ws, main_title, subtitle, table_title):
        ws['A1'] = f"Title: {main_title}"
        ws['A2'] = f"Subtitle: {subtitle}"
        ws['A3'] = f"Table Title: {table_title}"
        ws['A4'] = ""
        start_row = 5
        for r, table_row in enumerate(table.rows):
            for c, cell in enumerate(table_row.cells):
                ws.cell(row=start_row + r, column=c + 1, value=cell.text.strip())
        self.apply_enhanced_formatting(ws, start_row, len(table.rows))

    def apply_enhanced_formatting(self, ws, start_row, num_rows):
        title_fill = PatternFill(start_color='E6F3FF', end_color='E6F3FF', fill_type='solid')
        title_font = Font(bold=True, size=12)
        ws['A1'].font = title_font
        ws['A1'].fill = title_fill
        ws['A2'].font = Font(bold=True)
        ws['A3'].font = Font(bold=True)
        if num_rows > 0:
            header_fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
            for cell in ws[start_row]:
                if cell.value:
                    cell.font = Font(bold=True)
                    cell.fill = header_fill

    def create_summary_sheet(self, wb, filename, pages_converted, total_tables):
        ws = wb.create_sheet(title="Summary", index=0)
        ws['A1'] = "Selective PDF->DOCX->Excel Conversion Summary"
        ws['A1'].font = Font(size=16, bold=True)
        ws['A3'] = f"Source File: {filename}"
        ws['A4'] = f"Pages Converted: {pages_converted}"
        ws['A5'] = f"Total Tables Extracted: {total_tables}"
        ws['A6'] = f"Processing Method: Selective Page Conversion with Table Titles"

    def cleanup_temp_files(self):
        try:
            chunk_files = list(self.chunks_folder.glob("*.docx"))
            for chunk_file in chunk_files:
                chunk_file.unlink()
            logger.info(f"[INFO] Cleaned up {len(chunk_files)} temporary page files.")
        except Exception as e:
            logger.warning(f"Could not clean up temp files: {e}")


def find_all_folders_with_pdfs(root_path: Path, max_depth: int = 3) -> List[Tuple[Path, int, int]]:
    folders_with_pdfs = []
    def scan_directory(path: Path, current_depth: int = 0):
        if current_depth > max_depth: return
        try:
            pdf_files = list(path.glob("*.pdf"))
            if pdf_files:
                folders_with_pdfs.append((path, len(pdf_files), current_depth))
            for item in path.iterdir():
                if item.is_dir() and not item.name.startswith('.') and not item.name.startswith('__'):
                    scan_directory(item, current_depth + 1)
        except PermissionError: pass
    scan_directory(root_path)
    return folders_with_pdfs


def select_folder_from_current_directory():
    current_dir = Path.cwd()
    print(f"[SCAN] Scanning '{current_dir}' and subfolders for PDF files...")
    folders_with_pdfs = find_all_folders_with_pdfs(current_dir)
    if not folders_with_pdfs:
        print("[ERROR] No folders containing PDF files found!")
        return None
    folders_with_pdfs.sort(key=lambda x: (x[2], x[0].name.lower()))
    print(f"\nFound {len(folders_with_pdfs)} folders with PDF files:")
    print("-" * 80)
    for i, (folder_path, pdf_count, depth) in enumerate(folders_with_pdfs, 1):
        indent = "  " * depth
        relative_path = folder_path.relative_to(current_dir)
        display_path = f"{folder_path.name}" if depth == 0 else f"{indent}{relative_path}"
        print(f"{i:2d}. {display_path} ({pdf_count} PDF files)")
    print("-" * 80)
    while True:
        try:
            choice = input(f"\nSelect folder (1-{len(folders_with_pdfs)}): ").strip()
            if choice:
                folder_index = int(choice) - 1
                if 0 <= folder_index < len(folders_with_pdfs):
                    selected_folder, _, _ = folders_with_pdfs[folder_index]
                    print(f"[OK] Selected: {selected_folder.relative_to(current_dir)}")
                    return selected_folder
                else:
                    print(f"[ERROR] Please enter a number between 1 and {len(folders_with_pdfs)}")
            else:
                print("[ERROR] Please enter a folder number")
        except ValueError:
            print("[ERROR] Please enter a valid number")


def main():
    print("Title-Enhanced Selective Page PDF->DOCX->Excel Converter")
    print("=" * 75)
    print("This script will pre-scan PDFs and only convert pages with specific subtitles.")
    print("=" * 75)
    pdf_folder = select_folder_from_current_directory()
    if not pdf_folder:
        print("[ERROR] No folder selected. Exiting.")
        return
    excel_folder = Path.cwd() / f"{pdf_folder.name}_excel_output"
    print(f"--> Excel files will be saved to: {excel_folder}")
    max_workers = 6
    try:
        user_workers = input(f"\nMax workers (default {max_workers}): ").strip()
        if user_workers:
            max_workers = int(user_workers)
    except ValueError:
        print(f"Invalid input. Using default: {max_workers} workers.")
    converter = TitleEnhancedConverter(
        pdf_folder=str(pdf_folder),
        excel_output_folder=str(excel_folder), 
        max_workers=max_workers
    )
    converter.process_all_files()
    print(f"\n--> Intermediate DOCX files saved in: {converter.docx_folder}")
    print(f"--> Final Excel files saved in: {converter.excel_output_folder}")

if __name__ == "__main__":
    main()